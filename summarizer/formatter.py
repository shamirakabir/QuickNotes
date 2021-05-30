import docx
import random 
from datetime import date


class Formatter(object): 

    def __init__(self, lecture_file):
        self.lecture_file = lecture_file
        self.doc = docx.Document()

    def title(self):
        today = date.today()
        d1 = today.strftime("%d/%m/%Y")
        self.doc.add_heading('Introduction to Paging', 0)
        #header section
        header_section = self.doc.sections[0]
        header = header_section.header
        #header text
        header_text = header.paragraphs[0]
        header_text.text = "\t\t" + d1
    
    def bullet_list(self):
        self.title();
        count = 0;
        file = open(self.lecture_file, "r")
        first = False;
        second = False;
        for line in file:
            sentence = "";
            old = '';
            for c in line:
                if(c == '.'):
                    count += 1
                    #sentence += c
                    num = random.random()
                    if(num < 0.3 or count == 1):
                        self.doc.add_paragraph(sentence, style = 'List Bullet');
                        first = True;
                        second = False; 
                    elif(num > 0.3 and num <= 0.55):
                        if(first):
                            self.doc.add_paragraph(sentence, style = 'List Bullet 2');
                        else:
                            self.doc.add_paragraph(sentence, style = 'List Bullet');
                            second = True;
                    elif(num > 0.55 and num <= 1):
                        if(second):
                            self.doc.add_paragraph(sentence, style = 'List Bullet 3');
                            first = False;
                        elif(first):
                            self.doc.add_paragraph(sentence, style = 'List Bullet 2');
                            second = True;
                        else:
                            self.doc.add_paragraph(sentence, style = 'List Bullet');
                            second = False;
                    sentence = "";
                else:
                    sentence += c
    
    def save_docx(self):
        self.doc.save('temp.docx');

# def main():
#     format = Formatter('res.txt');
#     format.bullet_list();
#     format.save_docx();

if __name__ == "__main__":
    pass


# doc = docx.Document()

# doc.add_paragraph('The first item in an ordered list.',
#                   style='List Bullet')
  
# doc.add_paragraph('The second item in an ordered list.', 
#                   style='List Bullet')
  
# doc.add_paragraph('The third item in an ordered list.', 
#                   style='List Bullet')

# doc.save('temp.docx')
